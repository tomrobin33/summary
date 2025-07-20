from docx import Document

def parse_word(path):
    doc = Document(path)
    content = []
    for para in doc.paragraphs:
        if para.text.strip():
            content.append({"type": "paragraph", "text": para.text})
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            table_data.append([cell.text for cell in row.cells])
        content.append({"type": "table", "data": table_data})
    return content 